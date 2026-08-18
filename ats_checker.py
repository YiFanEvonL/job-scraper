# pip install python-docx
import sys
from docx import Document


def extract_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 ats_checker.py <resume.docx> \"keyword1,keyword2,...\"")
        sys.exit(1)

    docx_path = sys.argv[1]
    keywords = [k.strip() for k in sys.argv[2].split(",") if k.strip()]

    text = extract_text(docx_path).lower()

    found = []
    missing = []
    for kw in keywords:
        if kw.lower() in text:
            found.append(kw)
        else:
            missing.append(kw)

    for kw in found:
        print(f"  ✅  {kw}")
    for kw in missing:
        print(f"  ❌  {kw}")

    total = len(keywords)
    hits = len(found)
    pct = round(hits / total * 100) if total else 0
    print(f"\nATS 命中率：{hits}/{total} ({pct}%)")


if __name__ == "__main__":
    main()
